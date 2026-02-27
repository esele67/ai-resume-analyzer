from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS
import os, re, math, io, json
from collections import Counter

app = Flask(__name__, static_folder='.')
CORS(app)

# ─────────────────────────────────────────────────────────────
# OPTIONAL IMPORTS (graceful degradation)
# ─────────────────────────────────────────────────────────────
try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

try:
    import pypdf
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False

try:
    import docx as python_docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import pytesseract
    from pdf2image import convert_from_bytes
    HAS_OCR = True
except ImportError:
    HAS_OCR = False

try:
    import spacy
    nlp = spacy.load("en_core_web_sm")
    HAS_SPACY = True
except Exception:
    HAS_SPACY = False

try:
    import nltk
    for pkg in ['punkt', 'stopwords']:
        try:
            nltk.data.find(f'tokenizers/{pkg}')
        except LookupError:
            nltk.download(pkg, quiet=True)
    from nltk.corpus import stopwords as nltk_stopwords
    NLTK_STOPS = set(nltk_stopwords.words('english'))
    HAS_NLTK = True
except Exception:
    HAS_NLTK = False
    NLTK_STOPS = set()

# ─────────────────────────────────────────────────────────────
# CONSTANTS
# ─────────────────────────────────────────────────────────────
FALLBACK_STOPS = set("""
a an and are as at be by for from has he in is it its of on that the to was were will with
this but they have had what said each which she do how their if up out many then them these
so some her would make like into him time two more very when come may only see first way been
find use get made long little after words just where most know good much your work life day
new years year also still should could through being does now before here right too any about
i we you he she it they us our my his its who whom whose which what when where why how all
both few other some such no nor not only own same than then there can ive youve weve theyve
id youd hed shed wed ill youll hell shell well isnt arent wasnt werent havent hasnt hadnt
wont wouldnt dont doesnt didnt cant couldnt shouldnt
""".split())

STOP_WORDS = (NLTK_STOPS | FALLBACK_STOPS) if HAS_NLTK else FALLBACK_STOPS

HARD_SKILLS = {
    "Programming Languages": [
        "python","javascript","java","c++","c#","php","ruby","swift","kotlin","go","golang",
        "rust","typescript","r","matlab","scala","perl","lua","haskell","bash","shell",
        "html","css","sass","less","sql","nosql","assembly","cobol","fortran","dart","elixir"
    ],
    "Frameworks & Libraries": [
        "react","angular","vue","svelte","nextjs","next.js","nuxt","gatsby","nodejs","node.js",
        "express","fastapi","django","flask","spring","hibernate","tensorflow","pytorch","keras",
        "scikit-learn","sklearn","pandas","numpy","matplotlib","seaborn","plotly","jquery",
        "bootstrap","tailwind","laravel","rails","asp.net","flutter","graphql","redux","webpack"
    ],
    "Databases": [
        "mysql","postgresql","postgres","mongodb","sqlite","oracle","redis","cassandra",
        "elasticsearch","neo4j","firebase","dynamodb","mariadb","sql server","mssql",
        "couchdb","influxdb","snowflake","bigquery","hbase","supabase","planetscale"
    ],
    "Cloud & DevOps": [
        "aws","amazon web services","azure","gcp","google cloud","docker","kubernetes","k8s",
        "jenkins","git","github","gitlab","bitbucket","terraform","ansible","puppet","chef",
        "nginx","apache","ci/cd","devops","linux","unix","helm","prometheus","grafana",
        "datadog","cloudformation","circleci","github actions","argocd"
    ],
    "Data & ML": [
        "machine learning","deep learning","nlp","natural language processing","computer vision",
        "data science","data analysis","data engineering","big data","spark","hadoop","kafka",
        "airflow","tableau","power bi","looker","dbt","statistics","neural network","transformer",
        "bert","llm","reinforcement learning","feature engineering","mlops","a/b testing"
    ],
    "Security": [
        "cybersecurity","penetration testing","ethical hacking","owasp","encryption","ssl","tls",
        "siem","soc","vulnerability assessment","incident response","zero trust","iam","oauth","sso"
    ],
    "Certifications": [
        "pmp","aws certified","azure certified","google certified","comptia","cissp","ceh",
        "ccna","ccnp","itil","prince2","six sigma","cpa","cfa","scrum master","csm","safe","togaf"
    ]
}

SOFT_SKILLS = [
    "leadership","communication","teamwork","problem solving","critical thinking","creativity",
    "adaptability","time management","project management","analytical","collaboration","mentoring",
    "coaching","negotiation","presentation","conflict resolution","decision making",
    "emotional intelligence","attention to detail","organization","multitasking","self-motivated",
    "proactive","initiative","agile","scrum","cross-functional","stakeholder management"
]

SECTION_PATTERNS = {
    "contact":      [r'contact', r'email', r'phone', r'address', r'linkedin', r'github'],
    "summary":      [r'summary', r'objective', r'profile', r'about me', r'professional summary'],
    "experience":   [r'experience', r'employment', r'work history', r'professional experience', r'career history'],
    "education":    [r'education', r'academic', r'qualifications', r'degrees', r'university', r'college'],
    "skills":       [r'skills', r'technical skills', r'competencies', r'technologies', r'expertise'],
    "projects":     [r'projects', r'portfolio', r'work samples', r'personal projects'],
    "certifications":[r'certifications', r'certificates', r'licenses', r'credentials'],
    "achievements": [r'achievements', r'accomplishments', r'awards', r'honors'],
    "publications": [r'publications', r'papers', r'research'],
}

QUANTIFIABLE_PATTERNS = [
    r'\d+\s*%', r'\$\s*\d+', r'\d+\s*(?:million|billion|thousand|k\b)',
    r'increased\s+(?:by\s+)?\d+', r'reduced\s+(?:by\s+)?\d+',
    r'improved\s+(?:by\s+)?\d+', r'saved\s+(?:\$\s*)?\d+',
    r'\d+\s*(?:users|customers|clients|employees|team members)',
    r'\d+\s*(?:projects|products|features|applications)',
    r'\d+x\s*(?:improvement|faster|increase)', r'top\s*\d+',
]

WEAK_PHRASES = [
    "responsible for","worked on","helped with","assisted in","involved in",
    "participated in","contributed to","various","several","many tasks",
    "hard worker","team player","detail oriented","good communication",
    "fast learner","go getter","results driven","dynamic","synergy","leverage"
]

ACTION_VERBS = [
    "achieved","accelerated","architected","automated","built","championed","created",
    "delivered","designed","developed","directed","drove","engineered","established",
    "executed","founded","generated","grew","implemented","improved","increased",
    "initiated","launched","led","managed","mentored","modernized","optimized",
    "orchestrated","pioneered","reduced","refactored","scaled","shipped","solved",
    "spearheaded","streamlined","transformed","unified","upgraded"
]


# ─────────────────────────────────────────────────────────────
# TEXT EXTRACTION
# ─────────────────────────────────────────────────────────────
def extract_text_from_file(file):
    filename = file.filename.lower()
    content = file.read()

    if filename.endswith('.txt'):
        return content.decode('utf-8', errors='ignore'), False

    elif filename.endswith('.pdf'):
        if HAS_PDFPLUMBER:
            try:
                with pdfplumber.open(io.BytesIO(content)) as pdf:
                    text = '\n'.join(p.extract_text() or '' for p in pdf.pages).strip()
                    if len(text) > 100:
                        return text, False
            except Exception:
                pass
        if HAS_PYPDF:
            try:
                reader = pypdf.PdfReader(io.BytesIO(content))
                text = '\n'.join(p.extract_text() or '' for p in reader.pages).strip()
                if len(text) > 100:
                    return text, False
            except Exception:
                pass
        if HAS_OCR:
            try:
                images = convert_from_bytes(content, dpi=200)
                ocr_text = '\n'.join(pytesseract.image_to_string(img) for img in images)
                return ocr_text.strip(), True
            except Exception:
                pass
        return content.decode('latin-1', errors='ignore'), False

    elif filename.endswith('.docx'):
        if HAS_DOCX:
            try:
                doc = python_docx.Document(io.BytesIO(content))
                paragraphs = [p.text for p in doc.paragraphs]
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            paragraphs.append(cell.text)
                return '\n'.join(paragraphs), False
            except Exception:
                pass
        return content.decode('utf-8', errors='ignore'), False

    return content.decode('utf-8', errors='ignore'), False


# ─────────────────────────────────────────────────────────────
# NLP CORE
# ─────────────────────────────────────────────────────────────
def clean_text(text):
    text = text.lower()
    text = re.sub(r'[^\w\s\.\+\#\/]', ' ', text)
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

def tokenize(text):
    return [w for w in clean_text(text).split() if len(w) > 2 and w not in STOP_WORDS and not w.isdigit()]

def extract_keywords(text, top_k=40):
    freq = Counter(tokenize(text))
    return [{"word": w, "frequency": f} for w, f in freq.most_common(top_k)]

def tfidf_keywords(resume_text, jd_text, top_k=40):
    resume_freq = Counter(tokenize(resume_text))
    jd_words = tokenize(jd_text)
    total_jd = len(jd_words) or 1
    jd_freq = Counter(jd_words)
    results = {}
    for word, count in jd_freq.items():
        tf = count / total_jd
        n_docs = (1 if word in resume_freq else 0) + 1
        idf = math.log((3) / (n_docs + 0.5))
        results[word] = round(tf * idf, 5)
    return sorted(results.items(), key=lambda x: x[1], reverse=True)[:top_k]

def cosine_similarity(text1, text2):
    f1, f2 = Counter(tokenize(text1)), Counter(tokenize(text2))
    if not f1 or not f2:
        return 0.0
    vocab = set(f1) | set(f2)
    dot = sum(f1[w] * f2[w] for w in vocab)
    n1 = math.sqrt(sum(v**2 for v in f1.values()))
    n2 = math.sqrt(sum(v**2 for v in f2.values()))
    return dot / (n1 * n2) if n1 and n2 else 0.0


# ─────────────────────────────────────────────────────────────
# SKILLS
# ─────────────────────────────────────────────────────────────
def extract_skills(text):
    text_lower = text.lower()
    found, seen = [], set()
    for category, skill_list in HARD_SKILLS.items():
        for skill in skill_list:
            sl = skill.lower()
            if sl in text_lower and sl not in seen:
                seen.add(sl)
                idx = text_lower.find(sl)
                ctx = text_lower[max(0,idx-100):idx+100]
                conf = 0.5
                if any(w in ctx for w in ['experience','years','expert','senior','proficient']): conf += 0.3
                if any(w in ctx for w in ['advanced','certified','specialist','lead']): conf += 0.2
                found.append({"name": skill, "category": category, "type": "hard", "confidence": min(round(conf,2),1.0)})
    for skill in SOFT_SKILLS:
        sl = skill.lower()
        if sl in text_lower and sl not in seen:
            seen.add(sl)
            found.append({"name": skill, "category": "Soft Skills", "type": "soft", "confidence": 0.7})
    return found

def get_skills_balance(skills):
    hard = [s for s in skills if s["type"] == "hard"]
    soft = [s for s in skills if s["type"] == "soft"]
    total = len(hard) + len(soft)
    if total == 0:
        return {"hard_count":0,"soft_count":0,"hard_pct":0,"soft_pct":0,"balance":"Unknown"}
    hp = round(len(hard)/total*100)
    sp = 100 - hp
    balance = "Too technical — add soft skills" if hp >= 80 else ("Too generic — add technical skills" if sp >= 70 else "Well balanced")
    return {"hard_count":len(hard),"soft_count":len(soft),"hard_pct":hp,"soft_pct":sp,"balance":balance}


# ─────────────────────────────────────────────────────────────
# SECTION DETECTION & SCORING
# ─────────────────────────────────────────────────────────────
def detect_sections(text):
    lines = text.split('\n')
    sections, current, content = {}, "preamble", []
    for line in lines:
        ls = line.strip()
        matched = None
        for sname, pats in SECTION_PATTERNS.items():
            if any(re.search(p, ls, re.IGNORECASE) for p in pats) and len(ls) < 60:
                matched = sname
                break
        if matched:
            if content: sections[current] = '\n'.join(content).strip()
            current, content = matched, []
        else:
            content.append(ls)
    if content: sections[current] = '\n'.join(content).strip()
    return sections

def score_sections(sections, jd_text):
    jd_words = set(tokenize(jd_text))
    scores = {}
    for name, content in sections.items():
        if not content or name == "preamble": continue
        words = set(tokenize(content))
        if not words: continue
        overlap = words & jd_words
        score = round(min(len(overlap) / max(len(words)*0.3, 1), 1.0) * 100)
        scores[name] = {"score": score, "word_count": len(content.split()),
                        "matched_keywords": list(overlap)[:8], "present": True}
    for sec in ["experience","education","skills"]:
        if sec not in scores:
            scores[sec] = {"score": 0, "word_count": 0, "matched_keywords": [], "present": False}
    return scores


# ─────────────────────────────────────────────────────────────
# SPACY NER
# ─────────────────────────────────────────────────────────────
def extract_entities(text):
    base = {"organizations":[],"dates":[],"locations":[],"degrees":[],"titles":[]}
    if HAS_SPACY:
        doc = nlp(text[:10000])
        for ent in doc.ents:
            if ent.label_ == "ORG": base["organizations"].append(ent.text)
            elif ent.label_ == "DATE": base["dates"].append(ent.text)
            elif ent.label_ in ("GPE","LOC"): base["locations"].append(ent.text)
    degree_pat = r'\b(?:B\.?S\.?|B\.?A\.?|M\.?S\.?|M\.?A\.?|Ph\.?D\.?|MBA|BSc|MSc|BEng|MEng|B\.Tech|M\.Tech|Bachelor|Master|Doctorate)s?\b'
    base["degrees"] = list(dict.fromkeys(re.findall(degree_pat, text, re.IGNORECASE)))[:6]
    for key in base:
        base[key] = list(dict.fromkeys(base[key]))[:10]
    return base


# ─────────────────────────────────────────────────────────────
# CONTENT QUALITY
# ─────────────────────────────────────────────────────────────
def analyze_content_quality(text):
    sentences = [s.strip() for s in re.split(r'[.\n]+', text) if len(s.strip()) > 15]
    quantified, unquantified = [], []
    for sent in sentences:
        has_q = any(re.search(p, sent, re.IGNORECASE) for p in QUANTIFIABLE_PATTERNS)
        is_bullet = len(sent) < 180
        if has_q and is_bullet: quantified.append(sent[:130])
        elif is_bullet and not has_q and len(sent) > 25: unquantified.append(sent[:130])

    text_lower = text.lower()
    weak_found = [p for p in WEAK_PHRASES if p in text_lower]
    action_found = [v for v in ACTION_VERBS if re.search(r'\b'+v+r'\b', text_lower)]

    achievement_score = min(len(quantified) * 15, 100)
    action_score = min(len(action_found) * 6, 100)
    weak_penalty = min(len(weak_found) * 8, 40)
    content_score = max(round((achievement_score*0.5 + action_score*0.5) - weak_penalty), 0)

    return {
        "quantified_achievements": quantified[:6],
        "quantified_count": len(quantified),
        "unquantified_count": len(unquantified),
        "weak_phrases": weak_found[:8],
        "action_verbs_used": action_found[:12],
        "action_verb_count": len(action_found),
        "content_score": content_score,
        "achievement_score": achievement_score,
        "action_verb_score": action_score,
    }


# ─────────────────────────────────────────────────────────────
# FORMATTING
# ─────────────────────────────────────────────────────────────
def detect_formatting_issues(text):
    issues, score = [], 100
    sections = ['experience','education','skills']
    missing_secs = [s for s in sections if s not in text.lower()]
    if missing_secs:
        issues.append(f"Missing key sections: {', '.join(missing_secs)}")
        score -= 15
    wc = len(text.split())
    if wc < 150: issues.append(f"Too short ({wc} words) — aim for 300–700."); score -= 15
    elif wc > 1000: issues.append(f"Too long ({wc} words) — trim to 1–2 pages."); score -= 10
    has_email = bool(re.search(r'[\w.\-+]+@[\w.\-]+\.\w{2,}', text))
    if not has_email: issues.append("No email address detected."); score -= 10
    has_phone = bool(re.search(r'[\+\(]?\d[\d\s\-\(\)]{7,}\d', text))
    if not has_phone: issues.append("No phone number detected."); score -= 5
    has_bullets = bool(re.search(r'^[\s]*[•\-\*–]', text, re.MULTILINE))
    if not has_bullets: issues.append("No bullet points — hurts ATS readability."); score -= 8
    if re.search(r'[←→↑↓♦♣♠♥©®™€£¥§¶†‡◊№]', text): issues.append("Special characters found — ATS parsers reject these."); score -= 20
    has_linkedin = bool(re.search(r'linkedin\.com|linkedin', text, re.IGNORECASE))
    if not has_linkedin: issues.append("No LinkedIn URL — add your profile link."); score -= 5
    return {"score": max(score,0), "issues": issues, "word_count": wc,
            "has_email": has_email, "has_phone": has_phone, "has_bullets": has_bullets, "has_linkedin": has_linkedin}


# ─────────────────────────────────────────────────────────────
# HIGHLIGHTED PREVIEW
# ─────────────────────────────────────────────────────────────
def build_highlighted_preview(resume_text, matched_set, missing_set):
    lines = resume_text.split('\n')[:80]
    html_lines = []
    for line in lines:
        if not line.strip():
            html_lines.append('')
            continue
        words = line.split()
        highlighted = []
        for word in words:
            clean_w = re.sub(r'[^\w\+\#]', '', word.lower())
            if clean_w in matched_set:
                highlighted.append(f'<mark class="kw-match">{word}</mark>')
            elif clean_w in missing_set:
                highlighted.append(f'<mark class="kw-missing">{word}</mark>')
            else:
                highlighted.append(word)
        html_lines.append(' '.join(highlighted))
    return '<br>'.join(html_lines)


# ─────────────────────────────────────────────────────────────
# EXPERIENCE
# ─────────────────────────────────────────────────────────────
def extract_experience_years(text):
    patterns = [
        r'(\d{1,2})\+?\s*years?\s*(?:of\s*)?(?:professional\s*)?experience',
        r'experience\s*(?:of\s*)?(\d{1,2})\+?\s*years?',
        r'(\d{1,2})\+?\s*years?\s*(?:in|of)\s*(?:the\s*)?(?:industry|field|software|development)',
        r'(?:over|more\s*than)\s*(\d{1,2})\s*years?',
    ]
    years = 0
    for p in patterns:
        m = re.search(p, text, re.IGNORECASE)
        if m: years = max(years, int(m.group(1)))
    date_pat = r'(20\d{2}|19\d{2})\s*[-–—]\s*(20\d{2}|19\d{2}|present|current|now)'
    total_dates = 0
    for s, e in re.findall(date_pat, text, re.IGNORECASE):
        try:
            sy = int(s)
            ey = 2025 if e.lower() in ('present','current','now') else int(e)
            total_dates += max(0, ey - sy)
        except: pass
    return max(years, min(total_dates, 30))

def extract_required_experience(text):
    patterns = [
        r'(\d{1,2})\+?\s*years?\s*(?:of\s*)?(?:required|preferred|desired|minimum)',
        r'minimum\s*(?:of\s*)?(\d{1,2})\s*years?',
        r'at\s*least\s*(\d{1,2})\s*years?',
        r'(\d{1,2})\+?\s*years?\s*(?:of\s*)?experience\s*(?:required|preferred|needed)',
    ]
    years = 0
    for p in patterns:
        m = re.search(p, text, re.IGNORECASE)
        if m: years = max(years, int(m.group(1)))
    return years


# ─────────────────────────────────────────────────────────────
# NATURAL LANGUAGE SUMMARY
# ─────────────────────────────────────────────────────────────
def generate_nl_summary(overall, kw_pct, skills_pct, resume_exp, req_exp, content, missing_skills):
    lines = []
    if overall >= 80:
        lines.append(f"Your resume is strongly optimized for this role, achieving an ATS score of {overall}%.")
    elif overall >= 60:
        lines.append(f"Your resume scores {overall}% — a solid foundation with clear areas for improvement.")
    elif overall >= 40:
        lines.append(f"Your resume scores {overall}% — moderate alignment that needs targeted upgrades.")
    else:
        lines.append(f"Your resume scores {overall}% — significant optimization required to pass ATS screening.")
    lines.append(f"Keyword matching is at {kw_pct}% and skills alignment at {skills_pct}%.")
    if resume_exp > 0:
        exp_note = f"Your resume reflects approximately {resume_exp} year{'s' if resume_exp != 1 else ''} of experience"
        if req_exp > 0:
            exp_note += f" against a requirement of {req_exp}+ years"
            exp_note += " — you meet the threshold." if resume_exp >= req_exp else " — ensure all experience is visible."
        else:
            exp_note += "."
        lines.append(exp_note)
    if content["quantified_count"] > 3:
        lines.append(f"Strength: {content['quantified_count']} quantified achievements detected — this significantly boosts credibility.")
    else:
        lines.append(f"Weakness: only {content['quantified_count']} quantified achievement{'s' if content['quantified_count'] != 1 else ''} — add metrics to stand out.")
    if missing_skills:
        top3 = [s["name"] for s in missing_skills[:3]]
        lines.append(f"Key skill gaps: {', '.join(top3)}.")
    return " ".join(lines)


# ─────────────────────────────────────────────────────────────
# MAIN ANALYSIS ENGINE
# ─────────────────────────────────────────────────────────────
def analyze(resume_text, jd_text, resume_was_ocr=False):
    # Keywords
    jd_tfidf = tfidf_keywords(resume_text, jd_text, 40)
    jd_kw_list = [{"word": w, "score": s} for w, s in jd_tfidf]
    resume_kw = extract_keywords(resume_text, 50)
    resume_word_set = {k["word"] for k in resume_kw}
    jd_word_set = {k["word"] for k in jd_kw_list}

    matched_kw = [k for k in jd_kw_list if k["word"] in resume_word_set]
    missing_kw = [k for k in jd_kw_list if k["word"] not in resume_word_set]

    kw_match_pct = round(len(matched_kw) / max(len(jd_kw_list), 1) * 100)
    resume_tokens = tokenize(resume_text)
    total_tokens = len(resume_tokens) or 1
    kw_density = round(sum(1 for w in resume_tokens if w in jd_word_set) / total_tokens * 100, 1)

    # Skills
    resume_skills = extract_skills(resume_text)
    jd_skills = extract_skills(jd_text)
    jd_skill_names = {s["name"].lower() for s in jd_skills}
    resume_skill_names = {s["name"].lower() for s in resume_skills}

    matched_skills = [s for s in resume_skills if s["name"].lower() in jd_skill_names]
    missing_skills = [s for s in jd_skills if s["name"].lower() not in resume_skill_names]
    extra_skills = [s for s in resume_skills if s["name"].lower() not in jd_skill_names]
    skills_match_pct = round(len(matched_skills) / max(len(jd_skills), 1) * 100)
    skills_balance = get_skills_balance(resume_skills)

    by_cat = {}
    for s in resume_skills:
        by_cat[s["category"]] = by_cat.get(s["category"], 0) + 1

    # Experience
    resume_exp = extract_experience_years(resume_text)
    required_exp = extract_required_experience(jd_text)
    exp_match = min(resume_exp / max(required_exp, 1), 1.0) if required_exp else 0.7
    exp_level = ("Entry Level (0–2 yrs)" if resume_exp < 2 else
                 "Mid Level (2–5 yrs)" if resume_exp < 5 else
                 "Senior Level (5–10 yrs)" if resume_exp < 10 else "Executive Level (10+ yrs)")

    # Sections
    sections = detect_sections(resume_text)
    section_scores = score_sections(sections, jd_text)
    section_avg = round(sum(v["score"] for v in section_scores.values()) / max(len(section_scores), 1))

    # Formatting
    fmt = detect_formatting_issues(resume_text)

    # Content Quality
    content_quality = analyze_content_quality(resume_text)

    # NER
    entities = extract_entities(resume_text)

    # Cosine similarity
    cosine = cosine_similarity(resume_text, jd_text)

    # Highlighted preview
    matched_set = {k["word"] for k in matched_kw}
    missing_set = {k["word"] for k in missing_kw[:20]}
    highlighted_preview = build_highlighted_preview(resume_text, matched_set, missing_set)

    # Overall score
    weights = {"keyword":0.28,"skills":0.25,"experience":0.15,"formatting":0.10,
               "content":0.12,"sections":0.05,"cosine":0.05}
    overall = round(
        kw_match_pct * weights["keyword"] +
        skills_match_pct * weights["skills"] +
        exp_match * 100 * weights["experience"] +
        fmt["score"] * weights["formatting"] +
        content_quality["content_score"] * weights["content"] +
        section_avg * weights["sections"] +
        cosine * 100 * weights["cosine"]
    )
    overall = min(overall, 99)

    # Recommendations
    recs = []
    if resume_was_ocr:
        recs.append({"priority":"high","type":"formatting","title":"Scanned PDF Detected — Use Text-Based PDF",
            "description":"Your resume was read via OCR. Most ATS systems cannot process scanned PDFs. Export from Word/Google Docs/LaTeX as a text-based PDF."})
    if missing_kw:
        top_m = [k["word"] for k in missing_kw[:8]]
        recs.append({"priority":"high","type":"keywords","title":"Add Missing Keywords",
            "description":f"These JD keywords are absent from your resume: {', '.join(top_m)}. Add them naturally in your experience bullets and skills section."})
    if kw_density < 1.5:
        recs.append({"priority":"medium","type":"keywords","title":"Increase Keyword Density",
            "description":f"Keyword density is {kw_density}% — below the 2–4% ideal. Weave more role-specific terms naturally into your bullet points."})
    hard_missing = [s["name"] for s in missing_skills if s["type"]=="hard"][:5]
    soft_missing = [s["name"] for s in missing_skills if s["type"]=="soft"][:3]
    if hard_missing:
        recs.append({"priority":"high","type":"skills","title":"Add Required Technical Skills",
            "description":f"Required technical skills not found: {', '.join(hard_missing)}. Add to your Skills section if you have them."})
    if soft_missing:
        recs.append({"priority":"medium","type":"skills","title":"Demonstrate Soft Skills",
            "description":f"Show evidence of: {', '.join(soft_missing)} — through specific achievements, not just listing them."})
    if skills_balance["balance"] != "Well balanced":
        recs.append({"priority":"low","type":"skills","title":f"Skills Balance: {skills_balance['balance']}",
            "description":f"Your resume is {skills_balance['hard_pct']}% technical / {skills_balance['soft_pct']}% soft skills. Top candidates show both dimensions."})
    if required_exp > 0 and resume_exp < required_exp:
        recs.append({"priority":"high","type":"experience","title":"Experience Gap Detected",
            "description":f"Role requires {required_exp}+ years, resume shows ~{resume_exp}. List freelance, contract, and project experience with full date ranges."})
    if content_quality["quantified_count"] < 3:
        recs.append({"priority":"high","type":"content","title":"Quantify Your Achievements",
            "description":f"Only {content_quality['quantified_count']} quantified bullets found. Add metrics: 'Reduced latency by 40%', 'Managed 12-person team', 'Generated $500K revenue'."})
    if content_quality["weak_phrases"]:
        recs.append({"priority":"medium","type":"content","title":"Replace Weak Phrases",
            "description":f"Clichés detected: {', '.join(content_quality['weak_phrases'][:5])}. Replace with specific, measurable accomplishments."})
    if content_quality["action_verb_count"] < 5:
        recs.append({"priority":"medium","type":"content","title":"Use Strong Action Verbs",
            "description":f"Only {content_quality['action_verb_count']} power verbs found. Begin every bullet with: architected, spearheaded, scaled, delivered, optimized."})
    if not section_scores.get("projects",{}).get("present"):
        recs.append({"priority":"low","type":"sections","title":"Add a Projects Section",
            "description":"A Projects section with GitHub/demo links boosts technical ATS scores and showcases hands-on work."})
    if not fmt.get("has_linkedin"):
        recs.append({"priority":"low","type":"formatting","title":"Add LinkedIn Profile URL",
            "description":"87% of recruiters verify candidates on LinkedIn. Add your profile URL to the header."})
    for issue in fmt["issues"]:
        recs.append({"priority":"low","type":"formatting","title":"Formatting Issue","description":issue})

    recs.sort(key=lambda r: {"high":3,"medium":2,"low":1}.get(r["priority"],0), reverse=True)

    nl_summary = generate_nl_summary(overall, kw_match_pct, skills_match_pct,
                                     resume_exp, required_exp, content_quality, missing_skills)

    return {
        "overall_score": overall,
        "cosine_similarity": round(cosine * 100, 1),
        "was_ocr": resume_was_ocr,
        "capabilities": {"spacy": HAS_SPACY, "ocr": HAS_OCR, "nltk": HAS_NLTK},
        "nl_summary": nl_summary,
        "keyword_analysis": {
            "match_percentage": kw_match_pct, "keyword_density": kw_density,
            "matched_count": len(matched_kw), "missing_count": len(missing_kw),
            "total_jd_keywords": len(jd_kw_list),
            "matched_keywords": matched_kw[:15], "missing_keywords": missing_kw[:15],
        },
        "skills_analysis": {
            "match_percentage": skills_match_pct,
            "matched_count": len(matched_skills), "missing_count": len(missing_skills),
            "extra_count": len(extra_skills), "required_count": len(jd_skills),
            "matched_skills": matched_skills[:20], "missing_skills": missing_skills[:12],
            "extra_skills": extra_skills[:12], "by_category": by_cat,
            "balance": skills_balance,
        },
        "experience_analysis": {
            "detected_years": resume_exp, "required_years": required_exp,
            "match_score": round(exp_match * 100), "level": exp_level,
        },
        "section_analysis": section_scores,
        "content_quality": content_quality,
        "formatting_analysis": fmt,
        "entities": entities,
        "highlighted_preview": highlighted_preview,
        "recommendations": recs,
    }


# ─────────────────────────────────────────────────────────────
# BATCH ROUTE
# ─────────────────────────────────────────────────────────────
@app.route('/batch', methods=['POST'])
def batch_analyze():
    jd_text = request.form.get('jd_text', '')
    if 'jd_file' in request.files and request.files['jd_file'].filename:
        jd_text, _ = extract_text_from_file(request.files['jd_file'])
    if not jd_text.strip():
        return jsonify({"error": "Job description required."}), 400
    files = request.files.getlist('resume_files')
    if not files:
        return jsonify({"error": "No files uploaded."}), 400
    results = []
    for f in files[:10]:
        if not f.filename: continue
        try:
            text, was_ocr = extract_text_from_file(f)
            if not text.strip():
                results.append({"filename": f.filename, "error": "Could not extract text", "score": 0})
                continue
            r = analyze(text, jd_text, was_ocr)
            results.append({
                "filename": f.filename, "score": r["overall_score"],
                "keyword_match": r["keyword_analysis"]["match_percentage"],
                "skills_match": r["skills_analysis"]["match_percentage"],
                "experience_years": r["experience_analysis"]["detected_years"],
                "top_missing": [k["word"] for k in r["keyword_analysis"]["missing_keywords"][:5]],
                "nl_summary": r["nl_summary"]
            })
        except Exception as e:
            results.append({"filename": f.filename, "error": str(e), "score": 0})
    results.sort(key=lambda x: x.get("score", 0), reverse=True)
    return jsonify({"results": results, "count": len(results)})


# ─────────────────────────────────────────────────────────────
# ROUTES
# ─────────────────────────────────────────────────────────────
@app.route('/')
def index():
    return send_from_directory('.', 'web/index.html')

@app.route('/analyze', methods=['POST'])
def analyze_endpoint():
    resume_text, jd_text, was_ocr = '', '', False
    if 'resume_file' in request.files and request.files['resume_file'].filename:
        resume_text, was_ocr = extract_text_from_file(request.files['resume_file'])
    elif 'resume_text' in request.form:
        resume_text = request.form.get('resume_text', '')
    if 'jd_file' in request.files and request.files['jd_file'].filename:
        jd_text, _ = extract_text_from_file(request.files['jd_file'])
    elif 'jd_text' in request.form:
        jd_text = request.form.get('jd_text', '')
    if not resume_text.strip():
        return jsonify({"error": "Resume text is empty or could not be extracted."}), 400
    if not jd_text.strip():
        return jsonify({"error": "Job description is empty."}), 400
    if len(resume_text.strip()) < 50:
        return jsonify({"error": "Resume too short to analyze."}), 400
    try:
        return jsonify(analyze(resume_text, jd_text, was_ocr))
    except Exception as e:
        return jsonify({"error": f"Analysis failed: {str(e)}"}), 500

@app.route('/health')
def health():
    return jsonify({"status":"ok","capabilities":{
        "pdfplumber":HAS_PDFPLUMBER,"pypdf":HAS_PYPDF,"docx":HAS_DOCX,
        "ocr":HAS_OCR,"spacy":HAS_SPACY,"nltk":HAS_NLTK}})


if __name__ == '__main__':
    print("=" * 55)
    print("   ResumeIQ — Intelligent ATS Analyzer")
    print("   Open: http://localhost:5000")
    print("=" * 55)
    print(f"   pdfplumber : {'OK' if HAS_PDFPLUMBER else 'MISSING  pip install pdfplumber'}")
    print(f"   python-docx: {'OK' if HAS_DOCX else 'MISSING  pip install python-docx'}")
    print(f"   OCR support: {'OK' if HAS_OCR else 'MISSING  pip install pytesseract pdf2image'}")
    print(f"   spaCy NER  : {'OK' if HAS_SPACY else 'MISSING  pip install spacy && python -m spacy download en_core_web_sm'}")
    print(f"   NLTK       : {'OK' if HAS_NLTK else 'MISSING  pip install nltk'}")
    print("=" * 55)
    app.run(host="0.0.0.0", port=8080)